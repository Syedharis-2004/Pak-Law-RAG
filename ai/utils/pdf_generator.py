"""
PakLaw AI — PDF and DOCX Report Generators

Compiles AI legal research reports into professional PDFs (using ReportLab)
and DOCX files (using python-docx).
"""

import os
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document as WordDoc


def generate_pdf_report(report) -> str:
    """Generate a clean, printable PDF legal research report on disk."""
    output_dir = Path("./uploads/exports")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = output_dir / f"Report_{report.id}.pdf"
    
    doc = SimpleDocTemplate(str(file_path), pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom Styles
    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        spaceAfter=15
    )
    heading_style = ParagraphStyle(
        'ReportHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        spaceBefore=12,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        spaceAfter=8
    )

    story = []
    
    # Header
    story.append(Paragraph(report.title or "Legal Research Report", title_style))
    story.append(Spacer(1, 10))
    
    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Paragraph(report.executive_summary or "N/A", body_style))
    
    # Issues
    if report.legal_issues:
        story.append(Paragraph("Legal Issues & Analysis", heading_style))
        for item in report.legal_issues:
            issue_title = item.get('issue', 'Issue')
            analysis = item.get('analysis', 'Analysis')
            story.append(Paragraph(f"<b>Issue:</b> {issue_title}", body_style))
            story.append(Paragraph(f"<b>Analysis:</b> {analysis}", body_style))
            story.append(Spacer(1, 5))
            
    # Recommendations
    if report.recommendations:
        story.append(Paragraph("Practical Recommendations", heading_style))
        for rec in report.recommendations:
            story.append(Paragraph(f"• {rec}", body_style))
            
    doc.build(story)
    return str(file_path)


def generate_docx_report(report) -> str:
    """Generate a Microsoft Word document (.docx) of the legal research report."""
    output_dir = Path("./uploads/exports")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = output_dir / f"Report_{report.id}.docx"
    
    doc = WordDoc()
    
    # Title
    doc.add_heading(report.title or "Legal Research Report", 0)
    
    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.executive_summary or "N/A")
    
    # Issues
    if report.legal_issues:
        doc.add_heading("Legal Issues & Analysis", level=1)
        for item in report.legal_issues:
            p = doc.add_paragraph()
            p.add_run("Issue: ").bold = True
            p.add_run(item.get('issue', ''))
            
            p2 = doc.add_paragraph()
            p2.add_run("Analysis: ").bold = True
            p2.add_run(item.get('analysis', ''))
            
    # Recommendations
    if report.recommendations:
        doc.add_heading("Recommendations", level=1)
        for rec in report.recommendations:
            doc.add_paragraph(rec, style='List Bullet')
            
    doc.save(str(file_path))
    return str(file_path)
